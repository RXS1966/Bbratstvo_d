"""Генерация отчёта о проделанной работе (Variant_3) в формате Word."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Otchet-Neuroexam-Variant3.docx"


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def _add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def build_report() -> Document:
    doc = Document()
    _set_normal_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    today = date.today().strftime("%d.%m.%Y")

    _add_title(doc, "ОТЧЁТ О ПРОДЕЛАННОЙ РАБОТЕ")
    _add_paragraph(
        doc,
        "Проект: Нейроэкзаменатор HR MVP (Variant_3)",
    )
    p = doc.add_paragraph(f"Дата: {today}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    _add_heading(doc, "1. Цель работы")
    _add_paragraph(
        doc,
        "Разработать MVP системы оценки кандидатов в HR-процессе: "
        "диагностический срез, практические кейсы с автоматической "
        "оценкой ответов через LLM, итоговая сводка результатов и "
        "модуль экзаменатора. Обеспечить ролевой доступ (кандидат, "
        "руководитель, администратор), воспроизводимый запуск и "
        "документацию для передачи проекта.",
    )

    _add_heading(doc, "2. Архитектура решения")
    _add_paragraph(
        doc,
        "Реализован TypeScript monorepo на npm workspaces со следующей "
        "структурой:",
    )
    _add_bullet(doc, "packages/ui — общий UI-kit @repo/ui (shadcn-стиль)")
    _add_bullet(doc, "apps/web — фронтенд React + Vite + Tailwind (порт 5173)")
    _add_bullet(
        doc,
        "apps/api — бэкенд FastAPI + Alembic + SQLAlchemy (порт 8000)",
    )
    _add_bullet(
        doc,
        "PostgreSQL 16 в Docker (порт 5433 на хосте Windows)",
    )
    _add_paragraph(
        doc,
        "Фронтенд и API работают на одном домене: запросы /api "
        "проксируются на FastAPI. Авторизация — JWT, простой логин "
        "без SSO/LDAP.",
    )

    _add_heading(doc, "3. Выполненные задачи")
    tasks = [
        (
            "Разработана архитектура TypeScript monorepo "
            "«Нейроэкзаменатор HR MVP» на npm workspaces с разделением "
            "api/web/ui и базой данных в Docker PostgreSQL."
        ),
        (
            "Разработан FastAPI backend с ролями (кандидат, руководитель, "
            "admin), REST API для всех модулей HR-экзаменатора, миграциями "
            "Alembic и PostgreSQL."
        ),
        (
            "Разработан React-фронтенд (Vite, Tailwind) с UI-kit @repo/ui, "
            "7 экранами (логин, главная, диагностика, кейс, результат, "
            "экзаменатор, кабинет руководителя) и настраиваемой связкой "
            "с API."
        ),
        (
            "Внедрена LLM-оценка (OpenAI / Ollama) с промптами для кейсов, "
            "генерацией вопросов экзамена и пакетной (batch) проверкой "
            "ответов, включая llm_client с повторами и таймаутами."
        ),
        (
            "Реализован полный путь кандидата (диагностика → срез → кейс → "
            "результат → экзаменатор), ролевой доступ, демо-данные "
            "(seed_demo.py) и health-check (db_ok, llm_configured)."
        ),
        (
            "Настроены pytest (18 тестов, SQLite + mock LLM), Playwright "
            "E2E (7 сценариев) и CI в GitHub Actions (API, web-сборка, E2E)."
        ),
        (
            "Подготовлена документация (VNESHNIY-ZAPUSK.md, API-справочник, "
            "README, журнал roadmap A–E), архив neuroexam-variant3-mvp-*.zip "
            "и выполнен финальный прогон (pytest + build + E2E) перед "
            "передачей проекта."
        ),
    ]
    for task in tasks:
        _add_numbered(doc, task)

    _add_heading(doc, "4. Реализованный функционал")
    _add_paragraph(doc, "Экраны MVP и соответствующие API:")
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ("Раздел", "API", "Роли")
    rows_data = [
        ("Диагностика", "/api/diagnostic/sessions", "candidate, admin"),
        ("Кейс", "/api/cases", "candidate, admin"),
        ("Результат", "/api/results", "candidate, admin"),
        ("Экзаменатор", "/api/exam/sessions", "candidate, admin"),
        ("Кабинет руководителя", "/api/manager/overview", "manager, admin"),
    ]
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    for row_idx, row_data in enumerate(rows_data, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "Путь кандидата: диагностика → завершить срез → кейс → "
        "результат → экзаменатор. На главной странице отображается "
        "чеклист прогресса (CandidatePathChecklist).",
    )

    _add_heading(doc, "5. Технические улучшения (Roadmap A–E)")
    roadmap = [
        (
            "A — CI, health с БД, seed: проверка db_ok в /api/health, "
            "GitHub Actions (pytest + build), скрипт seed_demo.py."
        ),
        (
            "B — один LLM-вызов при завершении экзамена: пакетная оценка "
            "evaluate_exam_session_batch вместо трёх отдельных запросов."
        ),
        (
            "C — пользователи в PostgreSQL: таблица users, JWT-авторизация "
            "через БД, автосоздание demo/admin/manager при старте API."
        ),
        (
            "D — React Query и общие UI-компоненты: кэш health, мутации "
            "диагностики, WorkflowStatusBadge, BackendRequiredAlert."
        ),
        (
            "E — E2E Playwright и экспорт CSV для руководителя: "
            "GET /api/manager/export.csv, кнопка на ManagerPage."
        ),
    ]
    for item in roadmap:
        _add_bullet(doc, item)

    _add_heading(doc, "6. Тестирование и качество")
    _add_bullet(
        doc,
        "pytest: 18 автотестов API (diagnostic, cases, exam, results, "
        "health, manager, llm_client) на SQLite in-memory с моком LLM.",
    )
    _add_bullet(
        doc,
        "Playwright E2E: 7 сценариев — логин кандидата, полный путь MVP "
        "(диагностика, кейс, результат, экзаменатор), экспорт CSV manager.",
    )
    _add_bullet(
        doc,
        "CI (.github/workflows/ci.yml): job api (pytest), job web "
        "(npm ci + build), job e2e (Playwright с PostgreSQL service).",
    )
    _add_bullet(
        doc,
        "Health endpoint GET /api/health: status, db_ok, llm_configured.",
    )

    _add_heading(doc, "7. Документация и передача проекта")
    docs_items = [
        "README.md — обзор, запуск, хронология разработки (19 этапов).",
        "docs/VNESHNIY-ZAPUSK.md — пошаговый запуск для стороннего человека.",
        "docs/spravochnik-api-docs.md — справочник API ↔ экраны ↔ роли.",
        "docs/ROADMAP-ABCDE-ZHURNAL.md — журнал улучшений A–E.",
        "apps/api/README.md — миграции, seed, демо-логины.",
        "e2e/README.md — запуск Playwright E2E.",
        "Архив neuroexam-variant3-mvp-YYYYMMDD-roadmap-abcde.zip "
        "(скрипт Fronted/make_variant3_zip.py).",
    ]
    for item in docs_items:
        _add_bullet(doc, item)

    _add_heading(doc, "8. Демонстрация работы")
    _add_paragraph(doc, "Быстрый показ (около 5 минут):")
    demo_quick = [
        "Запустить Docker: docker compose up -d.",
        "В apps/api: миграции (.\\migrate.ps1), seed (python seed_demo.py), "
        "API (.\\run.ps1).",
        "В корне Variant_3: npm run dev, открыть http://localhost:5173/login.",
        "Войти как demo / demo — на главной видны бейджи API online и "
        "чеклист пути кандидата.",
        "Показать Результат (готовый срез и кейс с оценкой), Экзаменатор "
        "(черновик → старт → ответы → завершение с LLM-оценкой).",
        "Войти как manager / manager — кабинет руководителя, скачать CSV.",
    ]
    for step in demo_quick:
        _add_numbered(doc, step)

    _add_paragraph(doc, "")
    _add_paragraph(doc, "Демо-аккаунты:")
    _add_bullet(doc, "demo / demo — кандидат (полный путь)")
    _add_bullet(doc, "manager / manager — руководитель (обзор, CSV)")
    _add_bullet(doc, "admin / admin — администратор (все разделы)")

    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "Для LLM без OpenAI: Ollama (ollama pull llama3.2), в .env API "
        "указать OPENAI_BASE_URL=http://127.0.0.1:11434/v1, "
        "LLM_MODEL=llama3.2, OPENAI_API_KEY=ollama.",
    )

    _add_heading(doc, "9. Хронология разработки (основные этапы)")
    chronology = [
        "04.05.2026 — monorepo UI + web.",
        "18.05.2026 — FastAPI, UI-kit, диагностика, PostgreSQL + Alembic.",
        "18.05.2026 — экраны Кейс, Результат; роли и кабинет руководителя.",
        "18.05.2026 — LLM-оценка ответов на кейс.",
        "25.05.2026 — экран Экзаменатор, pytest, LLM-генерация вопросов.",
        "25.05.2026 — Roadmap A–E: CI, seed, batch LLM, users в БД, "
        "React Query, CSV export, Playwright E2E.",
    ]
    for item in chronology:
        _add_bullet(doc, item)

    _add_heading(doc, "10. Результат")
    _add_paragraph(
        doc,
        "Создано работоспособное MVP «Нейроэкзаменатор HR» с полным "
        "циклом оценки кандидата, интеграцией LLM, ролевым доступом, "
        "автотестами и документацией для воспроизводимого запуска. "
        "Проект готов к демонстрации, передаче коллеге и дальнейшему "
        "развитию (расширение промптов, новые роли, деплой).",
    )

    _add_paragraph(doc, "")
    _add_paragraph(doc, "Исполнитель: _________________________")
    _add_paragraph(doc, "Подпись: _________________________")

    return doc


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(OUTPUT)
    print(f"Отчёт сохранён: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
