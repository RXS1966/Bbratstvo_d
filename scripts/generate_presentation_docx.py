"""Генерация сценария презентации (Variant_3) в формате Word."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Stsenariy-prezentacii-Neuroexam.docx"


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def _add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def _add_table(
    doc: Document,
    headers: tuple[str, ...],
    rows: list[tuple[str, ...]],
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def _add_time_block(
    doc: Document,
    time_range: str,
    title: str,
    screen: str,
    actions: list[str],
    show: list[str],
    speech: str,
) -> None:
    _add_heading(doc, f"{time_range} — {title}")
    _add_paragraph(doc, "Что на экране:", bold=True)
    _add_paragraph(doc, screen)
    _add_paragraph(doc, "Действия:", bold=True)
    for action in actions:
        _add_bullet(doc, action)
    if show:
        _add_paragraph(doc, "Что показать:", bold=True)
        for item in show:
            _add_bullet(doc, item)
    _add_paragraph(doc, "Что сказать:", bold=True)
    _add_paragraph(doc, f"«{speech}»")


def build_presentation() -> Document:
    doc = Document()
    _set_normal_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    today = date.today().strftime("%d.%m.%Y")

    _add_title(doc, "СЦЕНАРИЙ ПРЕЗЕНТАЦИИ")
    _add_subtitle(doc, "Нейроэкзаменатор HR MVP (Variant_3)")
    _add_subtitle(doc, f"Дата: {today}")
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "Основной сценарий рассчитан на 12 минут с демо-данными (seed_demo.py). "
        "В конце документа — укороченный вариант (5 мин) и расширенный (20 мин).",
    )

    _add_heading(doc, "Подготовка (за 10–15 минут до начала)")
    prep_rows = [
        ("1", "docker compose up -d", "PostgreSQL в Docker"),
        ("2", "cd apps\\api → .\\migrate.ps1 → python seed_demo.py → .\\run.ps1", "API + демо-данные"),
        ("3", "В корне Variant_3: npm run dev", "Фронтенд на :5173"),
        ("4", "Ollama запущена, в apps/api/.env настроен LLM", "Бейдж LLM on"),
        ("5", "Открыть вкладки: 5173/login, 8000/api/docs, 8000/api/health", "Быстрый доступ"),
        ("6", "Закрыть лишние окна, масштаб браузера 110–125%", "Читаемость на проекторе"),
    ]
    _add_table(doc, ("№", "Команда / действие", "Назначение"), prep_rows)
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        "Проверка: после входа demo/demo на главной должны быть бейджи "
        "«API online» и «LLM on».",
        bold=True,
    )

    _add_heading(doc, "Основной сценарий — 12 минут")

    _add_time_block(
        doc,
        "0:00–0:45",
        "Вступление",
        "Слайд или титул проекта (можно сразу браузер).",
        [],
        [],
        "Добрый день. Покажу MVP «Нейроэкзаменатор HR» — систему оценки "
        "кандидатов: диагностический срез, практический кейс с оценкой через "
        "LLM, итоговая сводка и экзаменатор. Стек: React + FastAPI + "
        "PostgreSQL, monorepo с общим UI-kit.",
    )

    _add_time_block(
        doc,
        "0:45–1:30",
        "Вход и главная панель",
        "Страница http://localhost:5173/login → после входа блок «Панель».",
        [
            "Открыть http://localhost:5173/login",
            "Войти: demo / demo",
            "Убедиться, что открыта «Главная»",
        ],
        [
            "Бейдж «Кандидат»",
            "Бейджи «API online» и «LLM on»",
            "Чеклист «Ваш путь» — «Выполнено X из 5 шагов»",
        ],
        "Кандидат входит под demo. Система сразу показывает статус "
        "инфраструктуры: API, база, LLM. Чеклист ведёт по полному HR-пути — "
        "от диагностики до экзамена. Часть шагов уже выполнена демо-данными — "
        "это ускоряет показ.",
    )

    _add_time_block(
        doc,
        "1:30–3:00",
        "Готовые результаты (seed)",
        "Раздел «Результат» → детали среза «Руководитель продукта».",
        [
            "Меню → «Результат»",
            "Найти срез «Руководитель продукта»",
            "Нажать «Смотреть детали»",
        ],
        [
            "Статус среза: «завершён»",
            "Кейс «Приоритизация фич перед релизом» — статус «отправлен»",
            "Оценка 78/100 и текст feedback",
            "Контекст: B2B, KPI (MRR, конверсия, NPS)",
        ],
        "После диагностики кандидат решает кейс. Ответ уходит в LLM — "
        "возвращаются балл и развёрнутый feedback. Всё сохраняется в "
        "PostgreSQL. Руководитель видит структурированную сводку по срезу.",
    )
    _add_paragraph(
        doc,
        "Примечание: не читать весь feedback вслух — достаточно 1–2 предложений.",
    )

    _add_time_block(
        doc,
        "3:00–6:30",
        "Экзаменатор (главный акцент)",
        "Раздел «Экзаменатор» → карточка «Экзамен: Руководитель продукта».",
        [
            "Меню → «Экзаменатор»",
            "Карточка со статусом «черновик»",
            "«Запустить экзамен» (ждать 15–40 с с Ollama)",
            "Кратко ответить на 3 вопроса (по 1–2 предложения)",
            "«Завершить экзамен (LLM)» (ждать 30–60 с)",
        ],
        [
            "Статус «в работе» → «завершён»",
            "Итоговый балл (0–100)",
            "Feedback по каждому вопросу",
            "Текст «Экзамен завершён…»",
        ],
        "При старте LLM генерирует три вопроса с учётом роли и контекста "
        "среза. При завершении — один пакетный запрос: общий балл и оценка "
        "по каждому ответу. Это быстрее трёх отдельных вызовов.",
    )
    _add_paragraph(
        doc,
        "Запасной план: если LLM тормозит — заранее завершить один экзамен "
        "и показать готовый результат; live-запуск оставить как бонус.",
    )

    _add_time_block(
        doc,
        "6:30–7:30",
        "Возврат к сводке",
        "Раздел «Результат» → тот же срез → «Смотреть детали».",
        [
            "Меню → «Результат»",
            "Открыть детали среза «Руководитель продукта»",
        ],
        [
            "Появился блок экзамена со статусом «завершён»",
            "Единая картина: срез + кейс + экзамен",
        ],
        "Результат — единая точка: срез, кейсы и экзамены одного цикла "
        "оценки. Кандидат и руководитель видят полную картину.",
    )

    _add_time_block(
        doc,
        "7:30–9:00",
        "Кабинет руководителя",
        "Раздел «Кабинет руководителя» под учёткой manager.",
        [
            "Выйти или открыть инкогнито → manager / manager",
            "Меню → «Кабинет руководителя»",
            "Клик по кандидату demo (drill-down)",
            "«Скачать CSV» → открыть team-overview.csv",
        ],
        [
            "Список кандидатов и их срезы",
            "Детали: кейсы, экзамены кандидата demo",
            "Файл CSV для Excel",
        ],
        "Руководитель не создаёт срезы — только смотрит команду. Можно "
        "провалиться в кандидата demo. CSV — для отчётности без доработки "
        "интерфейса.",
    )

    _add_time_block(
        doc,
        "9:00–10:30",
        "Техническая часть",
        "Вкладки /api/health и /api/docs; опционально структура monorepo.",
        [
            "Показать http://localhost:8000/api/health (db_ok, llm_configured)",
            "Показать http://localhost:8000/api/docs (теги API)",
            "Опционально: apps/api, apps/web, packages/ui в проводнике",
        ],
        [
            "JSON health-check",
            "Swagger: diagnostic, cases, exam, manager",
            "Структура monorepo",
        ],
        "Health-check проверяет БД реальным запросом. Swagger — полный "
        "контракт API. Фронт на React Query, бэкенд FastAPI + Alembic. "
        "18 pytest-тестов, 7 Playwright-сценариев, CI в GitHub Actions. "
        "Запуск с нуля — docs/VNESHNIY-ZAPUSK.md.",
    )
    _add_paragraph(doc, "На демо pytest не запускать — только упомянуть.")

    _add_time_block(
        doc,
        "10:30–11:30",
        "Итог",
        "Главная или финальный слайд.",
        [],
        [],
        "Путь кандидата: диагностика → срез → кейс → результат → экзаменатор. "
        "Роли: demo — кандидат, manager — обзор, admin — всё. LLM — OpenAI "
        "или локальная Ollama. Итог: рабочий MVP с полным циклом оценки, "
        "автотестами и документацией. Готов к передаче и развитию.",
    )

    _add_time_block(
        doc,
        "11:30–12:00",
        "Вопросы",
        "Открыты: главная (demo), Swagger, CSV.",
        [
            "Держать готовые вкладки для быстрых ответов",
        ],
        [],
        "Готов ответить на вопросы по архитектуре, LLM, тестам и запуску.",
    )

    _add_heading(doc, "Типичные вопросы и короткие ответы")
    qa_rows = [
        ("Работает без интернета?", "Да, с локальной Ollama"),
        ("Где хранятся данные?", "PostgreSQL в Docker, порт 5433"),
        ("Как повторить запуск?", "docs/VNESHNIY-ZAPUSK.md"),
        ("Сколько тестов?", "18 pytest + 7 Playwright E2E"),
        ("Какие демо-логины?", "demo/demo, manager/manager, admin/admin"),
    ]
    _add_table(doc, ("Вопрос", "Ответ"), qa_rows)

    _add_heading(doc, "Укороченный сценарий — 5 минут")
    short_rows = [
        ("0:00", "Вступление (30 с)"),
        ("0:30", "demo/demo → Главная, бейджи + чеклист"),
        ("1:00", "Результат → срез с кейсом и оценкой"),
        ("2:00", "Экзаменатор → показать уже завершённый экзамен"),
        ("3:00", "manager → обзор + CSV"),
        ("4:00", "/api/health + фраза про тесты и документацию"),
        ("4:30", "Итог"),
    ]
    _add_table(doc, ("Минута", "Действие"), short_rows)

    _add_heading(doc, "Расширенный сценарий — 20 минут (без seed)")
    long_rows = [
        ("0:00–1:00", "Вступление + архитектура monorepo"),
        ("1:00–2:00", "Вход demo, главная, чеклист с 0 шагов"),
        ("2:00–5:00", "Диагностика: создать → запустить → завершить срез"),
        ("5:00–8:00", "Кейс: создать → ответ → «Отправить на оценку» (LLM)"),
        ("8:00–9:30", "Результат: сводка"),
        ("9:30–14:00", "Экзаменатор: создать → запустить → завершить (LLM)"),
        ("14:00–16:00", "manager + CSV"),
        ("16:00–18:00", "Swagger + health + CI"),
        ("18:00–20:00", "Итог + вопросы"),
    ]
    _add_table(doc, ("Минута", "Действие"), long_rows)
    _add_paragraph(
        doc,
        "Перед расширенным сценарием не запускать seed_demo.py "
        "или использовать чистую БД.",
    )

    _add_heading(doc, "Шпаргалка: если что-то сломалось")
    fallback_rows = [
        (
            "API offline",
            "Покажу Swagger / подготовленные данные",
            "Вкладка /api/docs",
        ),
        (
            "LLM off",
            "Оценка с ключом; покажу seed-оценку 78",
            "Раздел «Результат»",
        ),
        (
            "Долгий ответ LLM",
            "Локальная модель — 30–60 с, это нормально",
            "Заранее завершённый экзамен",
        ),
        (
            "Нет Docker",
            "Тесты API идут на SQLite",
            "pytest (по запросу)",
        ),
    ]
    _add_table(doc, ("Проблема", "Что сказать", "Запасной план"), fallback_rows)

    _add_heading(doc, "Чеклист в день презентации")
    checklist = [
        "seed_demo.py выполнен",
        "Один экзамен завершён заранее (запасной вариант)",
        "Ollama запущена (ollama list)",
        "3 вкладки: login, health, docs",
        "CSV скачан один раз (знать расположение файла)",
        "Масштаб экрана 110% и выше",
        "Демо-логины записаны на стикер: demo, manager, admin",
    ]
    for item in checklist:
        _add_bullet(doc, f"☐ {item}")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "Документ подготовлен для проекта Variant_3.")
    _add_paragraph(doc, "Связанные материалы: docs/VNESHNIY-ZAPUSK.md, Otchet-Neuroexam-Variant3.docx")

    return doc


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_presentation()
    doc.save(OUTPUT)
    print(f"Сценарий сохранён: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
